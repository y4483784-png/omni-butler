"""Document text extraction for Phase 2+.

Supported: .txt .md .csv .xlsx .pdf .docx
PDF low-text pages use Zhipu multimodal OCR when configured.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from app.core.config import settings
from app.core.messages import KB_PARSE_FAIL_MESSAGE

SUPPORTED = {
    ".txt",
    ".md",
    ".csv",
    ".xlsx",
    ".pdf",
    ".docx",
}


@dataclass
class ParsedElement:
    type: str
    text: str
    page: int | None = None
    heading: str = ""
    source: str = ""


@dataclass
class ParseResult:
    text: str
    elements: list[ParsedElement] | None = None
    warning: str = ""
    used_ocr: bool = False

    @property
    def char_count(self) -> int:
        return len(self.text or "")


_BLANKS_RE = re.compile(r"\n{3,}")
_SPACE_RUN_RE = re.compile(r" {2,}")
_NUMERIC_RE = re.compile(r"^-?\d+(?:\.\d+)?%?$")
_DATEISH_RE = re.compile(r"^\d{4}[-/年]\d{1,2}(?:[-/月]\d{1,2}日?)?$")


def parse_file(path: str | Path) -> ParseResult:
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"格式不支持：{ext or '(无扩展名)'}，支持 {', '.join(sorted(SUPPORTED))}")

    if ext in {".txt", ".md"}:
        return _read_text_elements(path)
    if ext == ".csv":
        return _read_csv(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    raise ValueError(f"格式不支持：{ext}")


def _display_stem(path: Path) -> str:
    stem = path.stem
    if "_" in stem and len(stem.split("_", 1)[0]) == 12:
        stem = stem.split("_", 1)[1]
    return stem


def _normalize_text(text: str) -> str:
    """Light cleanup for extracted prose while preserving structural markers."""
    if not text:
        return ""
    lines = [(_SPACE_RUN_RE.sub(" ", ln).rstrip()) for ln in text.splitlines()]
    text = "\n".join(lines)
    text = _BLANKS_RE.sub("\n\n", text)
    return text.strip()


def _make_element(
    type_: str,
    text: str,
    *,
    page: int | None = None,
    heading: str = "",
    source: str = "",
) -> ParsedElement | None:
    body = _normalize_text(text)
    if not body:
        return None
    return ParsedElement(type=type_, text=body, page=page, heading=heading, source=source)


def _render_elements(elements: list[ParsedElement]) -> str:
    return "\n\n".join(el.text for el in elements if (el.text or "").strip()).strip()


def _collect_non_empty_samples(rows: list[list[str]], col_idx: int, *, limit: int = 3) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for row in rows:
        if col_idx >= len(row):
            continue
        val = (row[col_idx] or "").strip()
        if not val or val in seen:
            continue
        seen.add(val)
        out.append(val[:24])
        if len(out) >= limit:
            break
    return out


def _tabular_hints(cols: list[str], rows: list[list[str]]) -> list[str]:
    if not cols:
        return []

    numeric_cols: list[str] = []
    date_cols: list[str] = []
    text_cols: list[str] = []
    sample_lines: list[str] = []

    probe_rows = rows[: min(len(rows), 50)]
    for idx, col in enumerate(cols[:20]):
        vals = [(r[idx] or "").strip() for r in probe_rows if idx < len(r) and (r[idx] or "").strip()]
        if vals:
            numeric_ratio = sum(1 for v in vals if _NUMERIC_RE.match(v)) / len(vals)
            date_ratio = sum(1 for v in vals if _DATEISH_RE.match(v)) / len(vals)
            if numeric_ratio >= 0.7:
                numeric_cols.append(col)
            elif date_ratio >= 0.6:
                date_cols.append(col)
            else:
                text_cols.append(col)

        samples = _collect_non_empty_samples(rows, idx)
        if samples:
            sample_lines.append(f"{col} 示例：{', '.join(samples)}")

    hints: list[str] = []
    if numeric_cols:
        hints.append("可能的数值列：" + "、".join(numeric_cols[:8]))
    if date_cols:
        hints.append("可能的时间列：" + "、".join(date_cols[:6]))
    if text_cols:
        hints.append("可能的维度列：" + "、".join(text_cols[:8]))
    hints.extend(sample_lines[:6])
    return hints


def _tabular_profile(path: Path, ext: str, cols: list[str], data_rows: int, rows: list[list[str]]) -> str:
    header = ", ".join(cols)
    lines = [
        f"数据集文件名：{_display_stem(path)}{ext}",
        f"列名（共 {len(cols)} 个）：{header}",
        f"数据行数：{data_rows}",
        f"数据列数：{len(cols)}",
    ]
    lines.extend(_tabular_hints(cols, rows))
    return "\n".join(lines)


def _tabular_sheet_context(sheet_name: str = "", title_note: str = "") -> str:
    lines: list[str] = []
    if sheet_name:
        lines.append(f"工作表：{sheet_name}")
    if title_note:
        lines.append(f"表格标题：{title_note}")
    return "\n".join(lines)


def _tabular_narrative_summary(
    path: Path,
    ext: str,
    cols: list[str],
    data_rows: int,
    rows: list[list[str]],
) -> str:
    name = f"{_display_stem(path)}{ext}"
    parts = [
        f"这是表格文件 {name} 的摘要。",
        f"共 {data_rows} 行、{len(cols)} 列。",
    ]
    if cols:
        parts.append("主要字段包括：" + "、".join(cols[:8]) + "。")

    hint_sentences: list[str] = []
    for hint in _tabular_hints(cols, rows)[:4]:
        if "：" in hint:
            label, value = hint.split("：", 1)
            hint_sentences.append(f"{label}有 {value}。")
        else:
            hint_sentences.append(hint + "。")
    parts.extend(hint_sentences)
    return "".join(parts)


def _format_row_batch(cols: list[str], start_row: int, rows: list[list[str]]) -> str:
    end_row = start_row + len(rows) - 1
    lines = [f"--- rows {start_row}-{end_row} ---", "\t".join(cols)]
    for cells in rows:
        lines.append("\t".join(cells))
    return "\n".join(lines)


def _format_row_batch_summary(cols: list[str], start_row: int, rows: list[list[str]]) -> str:
    end_row = start_row + len(rows) - 1
    samples: list[str] = []
    for row in rows[:2]:
        pairs = []
        for idx, col in enumerate(cols[:4]):
            if idx < len(row) and row[idx]:
                pairs.append(f"{col}={row[idx]}")
        if pairs:
            samples.append("；".join(pairs))
    sample_text = f" 示例记录：{' | '.join(samples)}。" if samples else ""
    return f"这是第 {start_row} 到 {end_row} 行的数据片段，字段围绕 {'、'.join(cols[:6])}。{sample_text}".strip()


def _assemble_tabular_text(
    path: Path,
    ext: str,
    cols: list[str],
    rows: list[list[str]],
    *,
    sheet_name: str = "",
    title_note: str = "",
) -> ParseResult:
    """Build profile + batched TSV blocks with row/char caps."""
    max_rows = settings.tabular_max_rows
    max_chars = settings.tabular_max_chars
    batch_size = max(1, settings.tabular_batch_rows)

    truncated_rows = len(rows) > max_rows
    if truncated_rows:
        rows = rows[:max_rows]

    profile = _tabular_profile(path, ext, cols, len(rows), rows)
    context = _tabular_sheet_context(sheet_name, title_note)
    if context:
        profile = f"{context}\n{profile}"
    summary = _tabular_narrative_summary(path, ext, cols, len(rows), rows)
    schema_text = _normalize_text(f"{summary}\n\n{profile}")
    parts: list[str] = [schema_text]
    elements: list[ParsedElement] = []
    source_name = f"tabular:{ext.lstrip('.')}" if not sheet_name else f"tabular:{ext.lstrip('.')}:{sheet_name}"
    summary_heading = "数据集摘要" if not sheet_name else f"{sheet_name} · 数据集摘要"
    row_heading_prefix = "" if not sheet_name else f"{sheet_name} · "
    schema_el = _make_element("schema", schema_text, heading=summary_heading, source=source_name)
    if schema_el:
        elements.append(schema_el)
    char_used = len(schema_text)
    truncated_chars = False

    for i in range(0, len(rows), batch_size):
        chunk = rows[i : i + batch_size]
        block = _normalize_text(f"{_format_row_batch_summary(cols, i + 1, chunk)}\n\n{_format_row_batch(cols, i + 1, chunk)}")
        extra = len(block) + 2
        if char_used + extra > max_chars:
            truncated_chars = True
            remain = max_chars - char_used - 2
            if remain > 80:
                truncated = block[:remain].rstrip()
                parts.append(truncated)
                row_el = _make_element(
                    "table",
                    truncated,
                    heading=f"{row_heading_prefix}rows {i + 1}-{i + len(chunk)}",
                    source=source_name,
                )
                if row_el:
                    elements.append(row_el)
            break
        parts.append(block)
        row_el = _make_element(
            "table",
            block,
            heading=f"{row_heading_prefix}rows {i + 1}-{i + len(chunk)}",
            source=source_name,
        )
        if row_el:
            elements.append(row_el)
        char_used += extra

    warnings: list[str] = []
    if truncated_rows:
        warnings.append(f"已截断至 {max_rows} 行（TABULAR_MAX_ROWS）")
    if truncated_chars:
        warnings.append(f"已截断至约 {max_chars} 字符（TABULAR_MAX_CHARS）")

    text = _normalize_text("\n\n".join(parts))
    return ParseResult(text=text, elements=elements, warning="；".join(warnings))


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")
    text = text.strip()
    if not text:
        raise ValueError("文件内容为空")
    return _normalize_text(text)


def _read_text_elements(path: Path) -> ParseResult:
    text = _read_text(path)
    source = "markdown" if path.suffix.lower() == ".md" else "text"
    elements: list[ParsedElement] = []
    heading = ""
    buf: list[str] = []

    def flush(kind: str = "text", *, use_heading: str = "") -> None:
        body = _normalize_text("\n".join(buf))
        if body:
            elements.append(
                ParsedElement(type=kind, text=body, heading=use_heading, source=source)
            )

    for line in text.splitlines():
        stripped = line.strip()
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            if buf:
                flush(use_heading=heading)
                buf = []
            heading = m.group(2).strip()
            elements.append(
                ParsedElement(type="heading", text=stripped, heading=heading, source=source)
            )
            continue
        if not stripped:
            if buf:
                flush(use_heading=heading)
                buf = []
            continue
        buf.append(line)
    if buf:
        flush(use_heading=heading)

    if not elements:
        fallback = _make_element("text", text, source=source)
        elements = [fallback] if fallback else []
    return ParseResult(text=text, elements=elements)


def _read_csv(path: Path) -> ParseResult:
    import csv
    from io import StringIO

    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "gbk", "latin-1"):
        try:
            decoded = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        decoded = raw.decode("utf-8", errors="replace")

    reader = csv.reader(StringIO(decoded))
    try:
        header = next(reader)
    except StopIteration:
        raise ValueError("CSV 内容为空") from None

    cols = [c.strip() for c in header if c.strip()]
    if not cols:
        raise ValueError("CSV 未解析到列名")

    max_rows = settings.tabular_max_rows
    rows: list[list[str]] = []
    truncated_rows = False
    for row in reader:
        cells = [c.strip() for c in row]
        if len(cells) < len(cols):
            cells.extend([""] * (len(cols) - len(cells)))
        cells = cells[: len(cols)]
        if not any(cells):
            continue
        rows.append(cells)
        if len(rows) >= max_rows:
            truncated_rows = next(reader, None) is not None
            break

    if not rows and not truncated_rows:
        # header-only is ok for schema questions
        pass

    result = _assemble_tabular_text(path, ".csv", cols, rows)
    if truncated_rows:
        extra = f"已截断至 {max_rows} 行（TABULAR_MAX_ROWS）"
        result.warning = extra if not result.warning else f"{result.warning}；{extra}"
    return result


def _cell_str(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _trim_trailing_empty(cells: list[str]) -> list[str]:
    out = list(cells)
    while out and not out[-1]:
        out.pop()
    return out


def _non_empty_count(cells: list[str]) -> int:
    return sum(1 for c in cells if c)


def _normalize_header_cells(row: list[str]) -> list[str]:
    trimmed = _trim_trailing_empty(row)
    cols: list[str] = []
    for idx, cell in enumerate(trimmed, start=1):
        cols.append(cell or f"column_{idx}")
    return cols


def _score_header_candidate(preview_rows: list[list[str]], idx: int) -> int:
    row = preview_rows[idx]
    non_empty = _non_empty_count(row)
    if non_empty == 0:
        return -10_000

    score = non_empty * 10
    if idx == 0 and non_empty == 1:
        score -= 25
    if idx > 0:
        score += 2

    for nxt in preview_rows[idx + 1 : idx + 4]:
        nxt_non_empty = _non_empty_count(nxt)
        if nxt_non_empty >= max(1, min(non_empty, 2)):
            score += 4
        if nxt_non_empty >= non_empty:
            score += 2
    return score


def _choose_header_row(preview_rows: list[list[str]]) -> int:
    candidates = [idx for idx, row in enumerate(preview_rows) if _non_empty_count(row) > 0]
    if not candidates:
        return 0

    wide_candidates = [idx for idx in candidates if _non_empty_count(preview_rows[idx]) >= 2]
    pool = wide_candidates or candidates
    return max(pool, key=lambda idx: (_score_header_candidate(preview_rows, idx), -idx))


def _prepare_sheet_table(ws) -> tuple[str, str, list[str], list[list[str]]]:
    rows_iter = ws.iter_rows(values_only=True)
    preview_rows: list[list[str]] = []
    for _ in range(12):
        raw = next(rows_iter, None)
        if raw is None:
            break
        preview_rows.append([_cell_str(c) for c in (raw or ())])

    if not preview_rows:
        return ws.title, "", [], []

    header_idx = _choose_header_row(preview_rows)
    title_parts = []
    for row in preview_rows[:header_idx]:
        non_empty = [c for c in row if c]
        if len(non_empty) == 1:
            title_parts.append(non_empty[0])
        elif non_empty:
            title_parts.append(" | ".join(non_empty[:4]))
    title_note = " / ".join(title_parts[:2])

    cols = _normalize_header_cells(preview_rows[header_idx])
    rows: list[list[str]] = []
    width = len(cols)
    for raw in preview_rows[header_idx + 1 :]:
        cells = _trim_trailing_empty(list(raw))
        if width:
            if len(cells) < width:
                cells.extend([""] * (width - len(cells)))
            cells = cells[:width]
        if any(cells):
            rows.append(cells)

    for raw in rows_iter:
        cells = _trim_trailing_empty([_cell_str(c) for c in (raw or ())])
        if width:
            if len(cells) < width:
                cells.extend([""] * (width - len(cells)))
            cells = cells[:width]
        if any(cells):
            rows.append(cells)

    return ws.title, title_note, cols, rows


def _read_xlsx(path: Path) -> ParseResult:
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise ValueError("缺少 openpyxl 依赖，无法解析 XLSX") from e

    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
    except Exception as e:
        raise ValueError(f"XLSX 打开失败：{e}") from e

    try:
        if not wb.sheetnames:
            raise ValueError("XLSX 中没有工作表")
        max_rows = settings.tabular_max_rows
        sheet_results: list[ParseResult] = []
        sheet_count = 0
        remaining_rows = max_rows
        truncated_rows = False
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_count += 1
            _, title_note, cols, rows = _prepare_sheet_table(ws)
            if not cols:
                continue

            if remaining_rows <= 0:
                truncated_rows = True
                break

            if len(rows) > remaining_rows:
                rows = rows[:remaining_rows]
                truncated_rows = True

            result = _assemble_tabular_text(
                path,
                ".xlsx",
                cols,
                rows,
                sheet_name=sheet_name,
                title_note=title_note,
            )
            sheet_results.append(result)
            remaining_rows -= len(rows)
    finally:
        wb.close()

    if not sheet_results:
        raise ValueError("XLSX 未解析到列名（请确认存在表头行）")

    all_text = "\n\n".join(r.text for r in sheet_results if r.text.strip())
    all_elements: list[ParsedElement] = []
    all_warnings: list[str] = []
    if sheet_count > 1:
        all_warnings.append(f"已解析 {len(sheet_results)}/{sheet_count} 个工作表")
    for r in sheet_results:
        all_elements.extend(r.elements or [])
        if r.warning:
            all_warnings.append(r.warning)

    result = ParseResult(
        text=_normalize_text(all_text),
        elements=all_elements,
        warning="；".join(w for w in all_warnings if w),
    )
    if truncated_rows:
        extra = f"已截断至 {max_rows} 行（TABULAR_MAX_ROWS）"
        result.warning = extra if not result.warning else f"{result.warning}；{extra}"
    return result


def _mime_from_content_type(content_type: str) -> str:
    ct = (content_type or "").lower()
    if "jpeg" in ct or "jpg" in ct:
        return "image/jpeg"
    if "png" in ct:
        return "image/png"
    if "webp" in ct:
        return "image/webp"
    if "gif" in ct:
        return "image/gif"
    if "bmp" in ct:
        return "image/bmp"
    return "image/png"


def _read_docx(path: Path) -> ParseResult:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ValueError("缺少 python-docx 依赖，无法解析 DOCX") from e

    from app.core.config import settings
    from app.rag.vision_ocr import ocr_image_bytes

    doc = DocxDocument(str(path))
    parts: list[str] = []
    elements: list[ParsedElement] = []
    active_heading = ""

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            style_name = (getattr(getattr(p, "style", None), "name", "") or "").lower()
            if style_name.startswith("heading"):
                active_heading = t
                el = _make_element("heading", t, heading=active_heading, source="docx")
            else:
                el = _make_element("text", t, heading=active_heading, source="docx")
            if el:
                elements.append(el)
            parts.append(t)

    for ti, table in enumerate(doc.tables, start=1):
        rows_md: list[str] = []
        header_cells: list[str] = []
        for row in table.rows:
            cells = [(" ".join((c.text or "").split())) for c in row.cells]
            if not header_cells and any(cells):
                header_cells = [c for c in cells if c]
            rows_md.append("| " + " | ".join(cells) + " |")
        if rows_md:
            header_text = "、".join(header_cells[:6]) if header_cells else "未识别到明确表头"
            summary = f"表格 {ti} 摘要：该表主要包含 {header_text} 等字段，共 {len(rows_md)} 行。"
            table_text = _normalize_text(f"[表格 {ti}]\n{summary}\n" + "\n".join(rows_md))
            parts.append(table_text)
            el = _make_element(
                "table",
                table_text,
                heading=active_heading or f"表格 {ti}",
                source="docx:table",
            )
            if el:
                elements.append(el)

    image_rels = []
    for rel in doc.part.rels.values():
        if "image" in (rel.reltype or ""):
            image_rels.append(rel)

    used_ocr = False
    ocr_ok = 0
    ocr_fail = 0
    warnings: list[str] = []
    ocr_ready = settings.vision_ocr_enabled and bool(settings.llm_api_key)

    if image_rels and not ocr_ready:
        warnings.append(
            f"文档含 {len(image_rels)} 张嵌入图，但未启用 OCR（需 LLM_API_KEY 且 VISION_OCR_ENABLED=true）"
        )

    for idx, rel in enumerate(image_rels, start=1):
        if not ocr_ready:
            break
        try:
            blob = rel.target_part.blob
            mime = _mime_from_content_type(getattr(rel.target_part, "content_type", "") or "")
            ocr_text = ocr_image_bytes(blob, mime=mime).strip()
            if ocr_text:
                ocr_block = _normalize_text(f"[嵌入图片 {idx} OCR]\n{ocr_text}")
                parts.append(ocr_block)
                el = _make_element(
                    "ocr",
                    ocr_block,
                    heading=active_heading or f"嵌入图片 {idx}",
                    source="docx:ocr",
                )
                if el:
                    elements.append(el)
                used_ocr = True
                ocr_ok += 1
            else:
                ocr_fail += 1
                warnings.append(f"嵌入图 {idx} OCR 无字")
        except Exception as e:
            ocr_fail += 1
            warnings.append(f"嵌入图 {idx} OCR 失败：{type(e).__name__}")

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("DOCX 中未找到文本、表格或可 OCR 的图片内容")

    if used_ocr:
        warnings.insert(0, f"已对 {ocr_ok}/{len(image_rels)} 张嵌入图使用多模态 OCR（{settings.vision_model}）")
    elif image_rels and ocr_fail:
        warnings.append(f"嵌入图 OCR 未产出文字（失败 {ocr_fail}）")

    return ParseResult(
        text=_normalize_text(text),
        elements=elements,
        warning="；".join(warnings),
        used_ocr=used_ocr,
    )


def _find_repeated_margin_lines(page_bodies: list[tuple[int, str]], *, min_pages: int = 3) -> tuple[set[str], set[str]]:
    top_counts: dict[str, int] = {}
    bottom_counts: dict[str, int] = {}
    if len(page_bodies) < min_pages:
        return set(), set()

    for _, body in page_bodies:
        lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
        if not lines:
            continue
        top = lines[0]
        bottom = lines[-1]
        if 1 <= len(top) <= 80:
            top_counts[top] = top_counts.get(top, 0) + 1
        if 1 <= len(bottom) <= 80:
            bottom_counts[bottom] = bottom_counts.get(bottom, 0) + 1

    threshold = max(min_pages, int(len(page_bodies) * 0.6))
    top_repeated = {ln for ln, cnt in top_counts.items() if cnt >= threshold}
    bottom_repeated = {ln for ln, cnt in bottom_counts.items() if cnt >= threshold}
    return top_repeated, bottom_repeated


def _strip_repeated_page_noise(page_bodies: list[tuple[int, str]]) -> tuple[list[tuple[int, str]], list[str]]:
    top_repeated, bottom_repeated = _find_repeated_margin_lines(page_bodies)
    if not top_repeated and not bottom_repeated:
        return page_bodies, []

    cleaned: list[tuple[int, str]] = []
    notes: list[str] = []
    for page_no, body in page_bodies:
        lines = [ln.rstrip() for ln in body.splitlines()]
        non_empty = [ln.strip() for ln in lines if ln.strip()]
        if not non_empty:
            cleaned.append((page_no, body))
            continue

        start = 0
        end = len(lines)
        while start < end and not lines[start].strip():
            start += 1
        while end > start and not lines[end - 1].strip():
            end -= 1

        if start < end and lines[start].strip() in top_repeated:
            start += 1
        if end > start and lines[end - 1].strip() in bottom_repeated:
            end -= 1

        cleaned.append((page_no, "\n".join(lines[start:end]).strip()))

    if top_repeated:
        notes.append(f"已去除重复页眉 {len(top_repeated)} 项")
    if bottom_repeated:
        notes.append(f"已去除重复页脚 {len(bottom_repeated)} 项")
    return cleaned, notes


def _read_pdf(path: Path) -> ParseResult:
    """Extract PDF text; OCR any page that has images (or almost no useful text)."""
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ValueError("缺少 pymupdf 依赖，无法解析 PDF") from e

    from app.rag.pdf_ocr_policy import collect_page_signals, decide_ocr, merge_native_and_ocr
    from app.rag.vision_ocr import ocr_image_bytes

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise ValueError(f"PDF 打开失败：{e}") from e

    warnings: list[str] = []
    used_ocr = False
    ocr_pages = 0
    skipped_cap = 0
    page_bodies: list[tuple[int, str]] = []
    max_ocr = settings.ocr_max_pages
    ocr_ready = settings.vision_ocr_enabled and bool(settings.llm_api_key)

    try:
        if getattr(doc, "is_encrypted", False):
            try:
                if not doc.authenticate(""):
                    raise ValueError(KB_PARSE_FAIL_MESSAGE)
            except Exception as e:
                raise ValueError(KB_PARSE_FAIL_MESSAGE) from e

        n_pages = doc.page_count
        for i in range(n_pages):
            page = doc.load_page(i)
            native = (page.get_text("text") or "").strip()
            signals = collect_page_signals(page, i)
            need, reason = decide_ocr(signals)
            page_text = native

            if need and ocr_ready:
                if ocr_pages >= max_ocr:
                    skipped_cap += 1
                else:
                    try:
                        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                        png = pix.tobytes("png")
                        ocr_text = ocr_image_bytes(png, mime="image/png").strip()
                        if ocr_text:
                            page_text = merge_native_and_ocr(native, ocr_text)
                            used_ocr = True
                            ocr_pages += 1
                        elif not page_text:
                            warnings.append(f"第{i + 1}页 OCR 无字（{reason}）")
                    except Exception as e:
                        warnings.append(f"第{i + 1}页 OCR 失败：{type(e).__name__}（{reason}）")
            elif need and not ocr_ready:
                warnings.append(
                    f"第{i + 1}页需 OCR（{reason}）但未启用：请配置 LLM_API_KEY 并保持 VISION_OCR_ENABLED=true"
                )

            if page_text:
                page_bodies.append((i + 1, _normalize_text(page_text)))
    finally:
        doc.close()

    page_bodies, noise_notes = _strip_repeated_page_noise(page_bodies)
    warnings.extend(noise_notes)
    elements: list[ParsedElement] = []
    parts = []
    for page_no, body in page_bodies:
        if not body:
            continue
        page_block = f"--- 第 {page_no} 页 ---\n{body}"
        parts.append(page_block)
        el = _make_element("page", page_block, page=page_no, heading=f"第 {page_no} 页", source="pdf")
        if el:
            elements.append(el)
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("未能从 PDF 提取到文本（可检查是否加密，或配置 LLM_API_KEY 启用多模态 OCR）")

    if used_ocr:
        warnings.insert(
            0,
            f"已对 {ocr_pages} 页使用多模态 OCR（有图即识别；模型 {settings.vision_model}）",
        )
    if skipped_cap:
        warnings.append(f"另有 {skipped_cap} 页因 OCR_MAX_PAGES={max_ocr} 未识别")
    if not used_ocr and not ocr_ready:
        # keep one concise hint if any page needed OCR
        pass
    elif len(text) < 200:
        warnings.append(f"抽出文字较少（{len(text)} 字），问答覆盖可能有限")

    # Dedupe nearly identical warning lines about missing key
    uniq: list[str] = []
    seen_w: set[str] = set()
    for w in warnings:
        key = w if "未启用" not in w else "ocr_disabled"
        if key in seen_w and key == "ocr_disabled":
            continue
        seen_w.add(key)
        uniq.append(w)

    return ParseResult(text=text, elements=elements, warning="；".join(uniq), used_ocr=used_ocr)
